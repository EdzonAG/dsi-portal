# modules/module5.py
from flask import Blueprint, render_template, request, flash, redirect, url_for, send_file, current_app
from flask_login import login_required
from models import module_permission_required, Module, PSPs, db
from datetime import date, datetime
from sqlalchemy import or_
from sqlalchemy.exc import IntegrityError
import unicodedata, io, re
import os
from docx import Document

# Palabras a ignorar al generar iniciales
STOPWORDS = {
    "de", "del", "la", "las", "el", "los", "y", "e", "en", "para", "por",
    "con", "a", "al", "da", "do", "das", "dos"
}

module5_bp = Blueprint('credenciales', __name__, template_folder='../templates')

BASE_VERIFY_URL = "https://toolbox.secomext.com.mx/credenciales/verify/"

# ---------------------- Utilidades ----------------------
def _parse_date(value: str):
    """Convierte 'YYYY-MM-DD' a datetime (00:00:00)."""
    try:
        return datetime.strptime(value, "%Y-%m-%d")
    except Exception:
        return None

def _strip_accents(s: str) -> str:
    return ''.join(c for c in unicodedata.normalize('NFD', s) if unicodedata.category(c) != 'Mn')

def _initials(s: str, *, skip_words=STOPWORDS) -> str:
    """
    Iniciales MAYÚSCULAS ignorando stopwords.
    'Dirección de Seguridad e Infraestructura' -> 'DSI'
    """
    s = _strip_accents(s or '')
    parts = re.split(r"[\s\-]+", s.strip().lower())
    letters = [p[0] for p in parts if p and p not in skip_words and p[0].isalpha()]
    return ''.join(letters).upper()

def _build_base_id(direccion: str, nombre: str) -> str:
    di = _initials(direccion)
    ni = _initials(nombre)
    return f"{di}-{ni}" if di and ni else (di or ni or "").upper()

def _unique_id(base_id: str) -> str:
    """Asegura unicidad: DSI-EOAG, DSI-EOAG-2, DSI-EOAG-3, ..."""
    candidate = base_id
    i = 2
    while PSPs.query.get(candidate):
        candidate = f"{base_id}-{i}"
        i += 1
    return candidate

def _creates_cycle(child_id: str, new_supervisor_id: str) -> bool:
    """
    Previene ciclos sencillos (p.ej. A->B y B->A).
    Recorre la cadena hacia arriba desde el nuevo supervisor.
    """
    cur = (new_supervisor_id or "").upper().strip()
    seen = set()
    while cur:
        if cur == child_id:
            return True
        if cur in seen:
            break
        seen.add(cur)
        sup = PSPs.query.get(cur)
        if not sup:
            break
        cur = (sup.supervisor_id or "").upper().strip()
    return False

def _fmt_ddmmyyyy(d):
    """Formatea date/datetime a DD/MM/YYYY."""
    if not d:
        return ""
    if isinstance(d, datetime):
        d = d.date()
    return d.strftime("%d/%m/%Y")

def _get_area_director_for(psp):
    """
    Director del área del PSP:
      - si el PSP tiene cadena jerárquica: su 'director' (penúltimo).
      - si es DG (no tiene jefe): él mismo.
    """
    try:
        director = psp.director()
    except Exception:
        director = None
    if director:
        return director.name
    return psp.name

def _find_director_for_area(area_hint):
    """
    Devuelve el director de un área por heurística.
    area_hint: "DO" o "DSI".
      - Busca un PSP con id que empiece por DO-/DSI- o por 'direccion' que contenga Operación / Seguridad e Infraestructura.
      - Devuelve su director (penúltimo). Si no hay, DG; si no, ese mismo PSP.
    """
    if area_hint == "DO":
        seed = PSPs.query.filter(
            or_(PSPs.id.ilike("DO-%"), PSPs.direccion.ilike("%Operación%"))
        ).first()
    else:  # DSI
        seed = PSPs.query.filter(
            or_(PSPs.id.ilike("DSI-%"), PSPs.direccion.ilike("%Seguridad e Infraestructura%"))
        ).first()

    if not seed:
        return None
    try:
        d = seed.get_director()
        return d or seed.get_director_general() or seed
    except Exception:
        return seed

def _replace_placeholder(doc, placeholder, replacement):
    replacement = "" if replacement is None else str(replacement)

    # Párrafos
    for p in doc.paragraphs:
        if placeholder in p.text:
            for run in p.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, replacement)

    # Tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if placeholder in p.text:
                        for run in p.runs:
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, replacement)

def _fill_docx_placeholders(input_path, data_dict):
    doc = Document(input_path)
    for ph, val in data_dict.items():
        _replace_placeholder(doc, ph, val)
    # guardar en memoria para descargar
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ---------------------- Listado / Home ----------------------
@module5_bp.route('/', methods=['GET'])
@login_required
@module_permission_required('credenciales')
def credenciales_home():
    search = (request.args.get('search') or '').strip()
    query = PSPs.query
    if search:
        like = f"%{search}%"
        query = query.filter(or_(
            PSPs.id.ilike(like),
            PSPs.name.ilike(like),
            PSPs.email.ilike(like),
            PSPs.direccion.ilike(like),
            PSPs.servicio.ilike(like),
            PSPs.telefono.ilike(like),
        ))
    history = query.order_by(PSPs.expiration_date).all()
    today = datetime.combine(date.today(), datetime.min.time())
    return render_template(
        'module5.html',
        show_sidebar=True,
        modules=Module.query.all(),
        history=history,
        search=search,
        module_name=Module.query.filter_by(identifier="credenciales").first().title,
        today=today
    )

# ---------------------- Crear PSP ----------------------
@module5_bp.route('/add', methods=['GET', 'POST'])
@login_required
@module_permission_required('credenciales')
def add_psp():
    if request.method == 'POST':
        name_psp   = (request.form.get('name_psp') or '').strip()
        email      = (request.form.get('email') or '').strip()
        direccion  = (request.form.get('direccion') or '').strip()
        servicio   = (request.form.get('servicio') or '').strip()
        telefono   = (request.form.get('telefono') or '').strip()
        status     = (request.form.get('status_psp') or 'inactivo').strip().lower()
        exp_str    = (request.form.get('expiration_date_psp') or '').strip()
        supervisor = (request.form.get('supervisor_id') or '').strip().upper() or None

        if not name_psp or not email or not direccion or not exp_str:
            flash("Nombre, correo, dirección y fecha de expiración son obligatorios.", "danger")
            return redirect(url_for('credenciales.add_psp'))

        exp_dt = _parse_date(exp_str)
        if not exp_dt:
            flash("Formato de fecha inválido. Usa YYYY-MM-DD.", "danger")
            return redirect(url_for('credenciales.add_psp'))

        # ID por defecto (único)
        base_id = _build_base_id(direccion, name_psp)
        if not base_id:
            flash("No se pudo generar el ID. Revisa nombre y dirección.", "danger")
            return redirect(url_for('credenciales.add_psp'))
        new_id = _unique_id(base_id)

        # Valida supervisor si viene
        if supervisor and not PSPs.query.get(supervisor):
            flash("El jefe seleccionado no existe. Se ignorará.", "warning")
            supervisor = None

        try:
            db.session.add(PSPs(
                id=new_id,
                name=name_psp,
                email=email,
                direccion=direccion,
                servicio=servicio,
                telefono=telefono,
                valid=(status == 'activo'),
                expiration_date=exp_dt,
                supervisor_id=supervisor
            ))
            db.session.commit()
            flash(f"PSP creado correctamente. ID asignado: {new_id}", "success")
            return redirect(url_for('credenciales.credenciales_home'))
        except IntegrityError:
            db.session.rollback()
            flash("No se pudo crear: ID duplicado.", "danger")
        except Exception as e:
            db.session.rollback()
            flash(f"Error al crear PSP: {e}", "danger")
        return redirect(url_for('credenciales.add_psp'))

    # GET: lista de posibles jefes
    all_psps = PSPs.query.order_by(PSPs.name).all()
    return render_template(
        'create_credentials.html',
        show_sidebar=True,
        modules=Module.query.all(),
        module_name=Module.query.filter_by(identifier="credenciales").first().title,
        all_psps=all_psps
    )

# ---------------------- Verificación pública ----------------------
@module5_bp.route('/verify/<psp_id>', methods=['GET'])
def verify_credentials(psp_id):
    psp = PSPs.query.get((psp_id or '').upper())
    if not psp:
        return render_template('verify_credentials.html', status="invalid", name=None)

    valid = bool(psp.valid)
    if psp.expiration_date and psp.expiration_date.date() < date.today():
        valid = False

    status = "ACTIVO" if valid else "INACTIVO"
    
    sup = psp.supervisor
    director = psp.director()
    
    return render_template(
        'verify_credentials.html',
        status=status,
        name=psp.name,
        psp_id=psp.id,
        direccion=psp.direccion,
        servicio=psp.servicio,
        correo=psp.email,
        jefe = director.name,
        expiration_date=psp.expiration_date
    )

# ---------------------- Editar PSP ----------------------
@module5_bp.route('/edit/<psp_id>', methods=['GET'])
@login_required
@module_permission_required('credenciales')
def edit_psp(psp_id):
    psp = PSPs.query.get((psp_id or '').upper())
    if not psp:
        flash("PSP no encontrado.", "danger")
        return redirect(url_for('credenciales.credenciales_home'))
    # Posibles jefes: todos menos él mismo
    all_psps = PSPs.query.filter(PSPs.id != psp.id).order_by(PSPs.name).all()
    return render_template(
        'edit_credentials.html',
        psp=psp,
        all_psps=all_psps,
        show_sidebar=True,
        modules=Module.query.all()
    )

@module5_bp.route('/update/<psp_id>', methods=['POST'])
@login_required
@module_permission_required('credenciales')
def update_credentials(psp_id):
    psp = PSPs.query.get((psp_id or '').upper())
    if not psp:
        flash("PSP no encontrado.", "danger")
        return redirect(url_for('credenciales.credenciales_home'))

    name_psp   = (request.form.get('name_psp') or '').strip()
    email      = (request.form.get('email') or '').strip()
    direccion  = (request.form.get('direccion') or '').strip()
    servicio   = (request.form.get('servicio') or '').strip()
    telefono   = (request.form.get('telefono') or '').strip()
    status     = (request.form.get('status_psp') or 'inactivo').strip().lower()
    exp_str    = (request.form.get('expiration_date_psp') or '').strip()
    supervisor = (request.form.get('supervisor_id') or '').strip().upper() or None

    if not name_psp or not email or not direccion or not exp_str:
        flash("Nombre, correo, dirección y fecha de expiración son obligatorios.", "danger")
        return redirect(url_for('credenciales.edit_psp', psp_id=psp.id))

    exp_dt = _parse_date(exp_str)
    if not exp_dt:
        flash("Formato de fecha inválido. Usa YYYY-MM-DD.", "danger")
        return redirect(url_for('credenciales.edit_psp', psp_id=psp.id))

    # Validaciones de jefe
    if supervisor == psp.id:
        flash("Un PSP no puede ser su propio jefe.", "danger")
        return redirect(url_for('credenciales.edit_psp', psp_id=psp.id))
    if supervisor and not PSPs.query.get(supervisor):
        flash("El jefe seleccionado no existe.", "danger")
        return redirect(url_for('credenciales.edit_psp', psp_id=psp.id))
    if supervisor and _creates_cycle(psp.id, supervisor):
        flash("Asignación inválida: generaría un ciclo de jerarquía.", "danger")
        return redirect(url_for('credenciales.edit_psp', psp_id=psp.id))

    try:
        psp.name = name_psp
        psp.email = email
        psp.direccion = direccion
        psp.servicio = servicio
        psp.telefono = telefono
        psp.valid = (status == 'activo')
        psp.expiration_date = exp_dt
        psp.supervisor_id = supervisor
        db.session.commit()
        flash("PSP actualizado correctamente.", "success")
        return redirect(url_for('credenciales.credenciales_home'))
    except Exception as e:
        db.session.rollback()
        flash(f"Error al actualizar PSP: {e}", "danger")
        return redirect(url_for('credenciales.edit_psp', psp_id=psp.id))

# ---------------------- Eliminar PSP ----------------------
@module5_bp.route('/delete/<psp_id>', methods=['POST'])
@login_required
@module_permission_required('credenciales')
def delete_credentials(psp_id):
    psp = PSPs.query.get((psp_id or '').upper())
    if psp:
        db.session.delete(psp)
        db.session.commit()
        flash("Registro eliminado", "success")
    else:
        flash("Registro no encontrado", "danger")
    return redirect(url_for('credenciales.credenciales_home'))

# ---------------------- QR descarga ----------------------
@module5_bp.route('/qr/<psp_id>', methods=['GET'])
@login_required
@module_permission_required('credenciales')
def qr_psp(psp_id):
    """Genera y descarga el QR con la URL pública de verificación."""
    import segno  # pip install segno
    url = f"{BASE_VERIFY_URL}{(psp_id or '').upper()}"
    qr = segno.make(url, error='m')
    buf = io.BytesIO()
    qr.save(buf, kind='png', scale=6)
    buf.seek(0)
    filename = f"{(psp_id or 'psp').upper()}_qr.png"
    return send_file(buf, mimetype='image/png', as_attachment=True, download_name=filename)


@module5_bp.route('/alta_docs/<psp_id>', methods=['GET'])
@login_required
@module_permission_required('credenciales')
def alta_docs(psp_id):
    # 1) Cargar PSP
    psp = PSPs.query.get((psp_id or "").upper())
    if not psp:
        flash("PSP no encontrado.", "danger")
        return redirect(url_for('credenciales.credenciales_home'))

    # 2) Resolver actores/fechas
    hoy = date.today()
    director = psp.director()  # director del área del PSP
    # RRHH: primer PSP cuyo servicio contenga 'RRHH'
    rrhh = PSPs.query.filter(PSPs.servicio.ilike("Coordinador de Procesos de Seguridad en RRHH")).first()

    # Director DSI/DO según prefijo del ID del PSP:
    #   - si empieza con DG-, DO- o C-  -> director de Operación (DO)
    #   - else                          -> director de Seguridad e Infraestructura (DSI)
    if psp.id.startswith(("DG-", "DO-", "C-")):
        dir_dsi_do = PSPs.query.filter(PSPs.servicio.ilike("Director de Operación")).first()
    else:
        dir_dsi_do = PSPs.query.filter(PSPs.servicio.ilike("Director de Seguridad e Infraestructura")).first()

    # 3) Construir mapping de placeholders
    data = {
        "[FECHA FIRMA]": _fmt_ddmmyyyy(hoy),
        "[NOMBRE PSP]": psp.name,
        "[REQUERIMIENTO PSP]": psp.servicio or "",
        "[FECHA INICIO]": _fmt_ddmmyyyy(hoy),
        "[FECHA FINAL]": _fmt_ddmmyyyy(psp.expiration_date),
        "[NOMBRE DIRECTOR]": director.name if director else psp.name,
        "[NOMBRE RRHH]": rrhh.name,
        "[DIRECTOR DSI/DO]": dir_dsi_do.name,
    }

    # 4) Plantilla y render
    #    Usa la carpeta instance de Flask: coloca el archivo dentro de /instance
    tpl_path = os.path.join(current_app.instance_path, "formats_templates/IFSW-N2-FOR_Servicios_Dist.docx")
    if not os.path.exists(tpl_path):
        flash("No se encontró la plantilla de Word en la carpeta 'instance'.", "danger")
        return redirect(url_for('credenciales.credenciales_home'))

    stream = _fill_docx_placeholders(tpl_path, data)

    # 5) Descargar
    filename = f"IFSW-N2-FOR_Servicios_Dist_{psp.id}.docx"
    return send_file(
        stream,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        as_attachment=True,
        download_name=filename
    )
